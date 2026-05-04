import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown

def shade_cell(cell, color):
    """Shade a cell with a color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def convert_md_to_docx():
    md_path = 'WinSCP_SFTP_Connection_Guide.md'
    docx_path = 'WinSCP_SFTP_Connection_Guide.docx'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run('WinSCP SFTP Connection Guide')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Set paragraph background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1e40af')
    title._element.get_or_add_pPr().append(shading_elm)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)
    
    # Add subtitle/metadata
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_data = [
        ('Protocol', 'SFTP (SSH File Transfer Protocol)'),
        ('Service', 'Azure Blob Storage'),
        ('Last Updated', 'April 13, 2026'),
        ('Status', 'Production')
    ]
    
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        shade_cell(row.cells[0], '1e40af')
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Parse markdown and add to document
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.style = 'Heading 1'
            i += 1
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.style = 'Heading 2'
            i += 1
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.style = 'Heading 3'
            i += 1
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.style = 'Heading 4'
            i += 1
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_block = doc.add_paragraph('\n'.join(code_lines), style='Intense Quote')
            i += 1
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Find end of table
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            # Parse table
            header = [col.strip() for col in table_lines[0].split('|')[1:-1]]
            rows = []
            for table_line in table_lines[2:]:
                cols = [col.strip() for col in table_line.split('|')[1:-1]]
                if cols and any(cols):
                    rows.append(cols)
            
            if header and rows:
                table = doc.add_table(rows=len(rows) + 1, cols=len(header))
                table.style = 'Light Grid Accent 1'
                
                # Header row
                for j, h in enumerate(header):
                    table.rows[0].cells[j].text = h
                    shade_cell(table.rows[0].cells[j], '1e40af')
                    row_cell = table.rows[0].cells[j].paragraphs[0]
                    if row_cell.runs:
                        row_cell.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        row_cell.runs[0].font.bold = True
                
                # Data rows
                for r, row_data in enumerate(rows, 1):
                    for c, col in enumerate(row_data):
                        if c < len(table.rows[r].cells):
                            table.rows[r].cells[c].text = col
        # Lists
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        # Regular paragraphs
        elif line.strip():
            p = doc.add_paragraph(line)
            i += 1
        else:
            i += 1
    
    # Save document
    doc.save(docx_path)
    print('✅ Word document created: WinSCP_SFTP_Connection_Guide.docx')

if __name__ == '__main__':
    convert_md_to_docx()
